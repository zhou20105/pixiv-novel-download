from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
import os
import re

def convert_cookies(input_file):
    """ 从文件中读取 cookies 并返回格式化后的列表 """
    with open(input_file, 'r', encoding='utf-8') as f:
        cookie_string = f.read().strip()
    
    cookies = []
    cookie_entries = cookie_string.split('; ')
    
    for entry in cookie_entries:
        name, value = entry.split('=', 1)
        cookies.append({
            "name": name,
            "value": value,
            "domain": ".pixiv.net",
            "path": "/"
        })
    return cookies

def sanitize_filename(filename):
    """去除文件名中的非法字符"""
    return re.sub(r'[\\/*?:"<>|]', "", filename)

def extract_novel_content(driver, novel_id):
    """根据小说ID提取小说标题和内容"""
    url = f"https://www.pixiv.net/novel/show.php?id={novel_id}"
    driver.get(url)
    
    try:
        # 等待标题元素加载
        title_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, 'sc-1u8nu73-3'))
        )
        novel_title = title_element.text.strip()

        # 等待内容元素加载
        content_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, 'sc-khIgEk'))
        )
        paragraphs = content_element.find_elements(By.TAG_NAME, 'p')

        return novel_title, paragraphs
    except Exception as e:
        print(f"提取小说内容时发生错误: {e}")
        return None, None

def save_novel_to_word(novel_title, paragraphs):
    """将提取的内容保存到Word文档"""
    doc = Document()
    doc.add_heading(novel_title, level=1)
    
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text:
            doc.add_paragraph(text)

    # 创建一个合法的文件名并保存文档
    sanitized_title = sanitize_filename(novel_title)
    file_name = f"{sanitized_title}.docx"
    save_path = os.path.join(os.getcwd(), file_name)
    doc.save(save_path)
    print(f"小说已保存到: {save_path}")

def setup_driver():
    """设置并返回Selenium WebDriver"""
    service = Service(ChromeDriverManager().install())
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')  # 无头模式
    options.add_argument('--disable-gpu')  # 禁用GPU加速
    options.add_argument('--no-sandbox')  # 防止沙盒问题
    
    driver = webdriver.Chrome(service=service, options=options)
    return driver

def load_cookies(driver, cookies):
    """加载cookies到浏览器"""
    driver.get("https://www.pixiv.net")
    for cookie in cookies:
        driver.add_cookie(cookie)

def main():
    # 从文件中读取cookies并格式化
    input_file = "cookie.txt"
    cookies = convert_cookies(input_file)
    
    # 设置driver并加载cookies
    driver = setup_driver()
    load_cookies(driver, cookies)
    
    while True:
        novel_id = input("请输入Pixiv小说ID（或输入 'exit' 退出）：").strip()
        
        if novel_id.lower() == 'exit':
            print("退出程序")
            driver.quit()
            break
        
        if not novel_id.isdigit():
            print("ID格式无效，请输入数字形式的小说ID")
            continue
        
        # 提取并保存小说内容
        novel_title, paragraphs = extract_novel_content(driver, novel_id)
        if novel_title and paragraphs:
            save_novel_to_word(novel_title, paragraphs)

if __name__ == "__main__":
    main()
