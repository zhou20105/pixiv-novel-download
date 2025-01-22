import os
import re
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document

def convert_cookies(input_file):
    """ 从文件中读取 cookies 并返回格式化后的列表 """
    with open(input_file, 'r', encoding='utf-8') as f:
        cookie_string = f.read().strip()
    
    cookies = []
    cookie_entries = cookie_string.split('; ')
    
    for entry in cookie_entries:
        name, value = entry.split('=', 1)
        cookies.append({
            "name": name,
            "value": value,
            "domain": ".pixiv.net",
            "path": "/"
        })
    return cookies

# 去除非法字符，防止保存出错
def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename)

# 提取网页标题和正文内容
def extract_novel_content(driver, url, chapter_number):
    driver.get(url)
    time.sleep(3)  # 等待页面加载完成
    
    try:
        # 获取章节标题
        title_element = driver.find_element(By.CLASS_NAME, 'sc-1u8nu73-3')
        novel_title = title_element.text.strip() if title_element else f"Chapter {chapter_number}"
        
        # 获取章节内容
        content_element = driver.find_element(By.CLASS_NAME, 'sc-khIgEk')
        paragraphs = content_element.find_elements(By.TAG_NAME, 'p')
        
        content = "\n".join([p.text.strip() for p in paragraphs if p.text.strip()])
        return novel_title, content
    
    except Exception as e:
        print(f"Error extracting content: {e}")
        return None, None

# 保存内容到Word文档，并加上章节序号
def save_to_word(novel_title, content, folder_path, chapter_number):
    doc = Document()
    doc.add_heading(f"第 {chapter_number} 章: {novel_title}", level=1)
    doc.add_paragraph(content)
    
    # 文件名合法化并加上章节序号
    file_name = f"第{chapter_number}章_{sanitize_filename(novel_title)}.docx"
    save_path = os.path.join(folder_path, file_name)
    doc.save(save_path)
    print(f"Saved to {save_path}")

# 提取系列章节URL
def extract_novel_urls(driver, series_url, cookies):
    driver.get("https://www.pixiv.net")
    time.sleep(2)  # 等待页面加载

    # 添加 cookies
    for cookie in cookies:
        driver.add_cookie(cookie)
    
    aurls = []
    for i in range(1, 101):  # 假设最多分页100
        driver.get(f"{series_url}?p={i}")
        time.sleep(3)
        
        # 获取章节链接
        novel_elements = driver.find_elements(By.CLASS_NAME, 'sc-1c4k3wn-12')
        if not novel_elements:
            break
        
        for novel_element in novel_elements:
            novel_url = novel_element.find_element(By.TAG_NAME, 'a')
            aurls.append(novel_url.get_attribute('href'))
    return aurls

# 获取系列名称并创建文件夹
def create_series_folder(driver, series_url, cookies):
    driver.get(series_url)
    time.sleep(3)
    
    try:
        # 获取系列标题
        series_element = driver.find_element(By.CLASS_NAME, 'sc-vk2fvc-2')
        series_title = sanitize_filename(series_element.text.strip()) if series_element else "Unknown Series"
        
        # 创建文件夹
        os.makedirs(series_title, exist_ok=True)
        return series_title
    except Exception as e:
        print(f"Error creating folder: {e}")
        return None

# 主逻辑
def main():
    input_file = "cookie.txt"
    cookies = convert_cookies(input_file)
    
    while True:
        series_id = input("Enter the Pixiv series ID (or 'exit' to quit): ").strip()
        if series_id.lower() == 'exit':
            print("Exiting the program.")
            break
        
        # 准备浏览器
        service = Service(ChromeDriverManager().install())
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        
        driver = webdriver.Chrome(service=service, options=options)
        
        try:
            series_url = f"https://www.pixiv.net/novel/series/{series_id}"
            
            # 获取系列名称和创建文件夹
            series_folder = create_series_folder(driver, series_url, cookies)
            if not series_folder:
                print("Failed to create series folder. Skipping.")
                continue
            
            # 获取章节URL
            novel_urls = extract_novel_urls(driver, series_url, cookies)
            
            # 提取并保存每一章
            for chapter_number, novel_url in enumerate(novel_urls, start=1):
                novel_title, content = extract_novel_content(driver, novel_url, chapter_number)
                if novel_title and content:
                    save_to_word(novel_title, content, series_folder, chapter_number)
        
        except Exception as e:
            print(f"An error occurred: {e}")
        
        finally:
            driver.quit()

if __name__ == "__main__":
    main()
