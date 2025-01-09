import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkinter import filedialog
import os
import re
import time
from concurrent.futures import ProcessPoolExecutor
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document


def create_driver(cookies):
    """ 创建独立的浏览器实例并加载 cookies """
    service = Service(ChromeDriverManager().install())
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')  # 无头模式
    options.add_argument('--disable-gpu')  # 禁用 GPU 加速
    options.add_argument('--no-sandbox')  # 防止沙盒问题
    driver = webdriver.Chrome(service=service, options=options)
    
    # 加载 cookies
    driver.get("https://www.pixiv.net")  # 必须访问页面才能添加 cookies
    for cookie in cookies:
        driver.add_cookie(cookie)  # 添加 cookies
    return driver


# 去除非法字符，防止保存出错
def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename)


# 提取网页标题和正文内容
def extract_novel_content(driver, url, chapter_number):
    driver.get(url)
    WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, 'sc-khIgEk')))  # 等待内容加载
    
    try:
        # 获取章节标题
        title_element = driver.find_element(By.CLASS_NAME, 'sc-1u8nu73-3')
        novel_title = title_element.text.strip() if title_element else f"Chapter {chapter_number}"
        
        # 获取章节内容
        content_element = driver.find_element(By.CLASS_NAME, 'sc-khIgEk')
        paragraphs = content_element.find_elements(By.TAG_NAME, 'p')
        
        content = "\\n".join([p.text.strip() for p in paragraphs if p.text.strip()])
        return novel_title, content
    
    except Exception as e:
        print(f"Error extracting content: {e}")
        return None, None


# 保存内容到Word文档，并加上章节序号
def save_to_word(novel_title, content, folder_path, chapter_number):
    doc = Document()
    doc.add_heading(f"第 {chapter_number} 章: {novel_title}", level=1)
    doc.add_paragraph(content)
    
    # 文件名合法化并加上章节序号
    file_name = f"第{chapter_number}章_{sanitize_filename(novel_title)}.docx"
    save_path = os.path.join(folder_path, file_name)
    doc.save(save_path)
    print(f"Saved to {save_path}")


# 提取系列章节URL
def extract_novel_urls(driver, series_url):
    aurls = []
    for i in range(1, 101):  # 假设最多分页100
        driver.get(f"{series_url}?p={i}")
        time.sleep(3)  # 等待章节加载
        
        # 获取章节链接
        novel_elements = driver.find_elements(By.CLASS_NAME, 'sc-1c4k3wn-12')
        if not novel_elements:
            break
        
        for novel_element in novel_elements:
            novel_url = novel_element.find_element(By.TAG_NAME, 'a')
            aurls.append(novel_url.get_attribute('href'))
    return aurls


# 获取系列名称并创建文件夹
def create_series_folder(driver, series_url, output_path=None):
    driver.get(series_url)
    WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, 'sc-vk2fvc-2')))  # 等待系列标题加载
    
    try:
        # 获取系列标题
        series_element = driver.find_element(By.CLASS_NAME, 'sc-vk2fvc-2')
        series_title = sanitize_filename(series_element.text.strip()) if series_element else "Unknown Series"
        
        # 创建文件夹
        if output_path:
            folder_path = os.path.join(output_path, series_title)
        else:
            folder_path = series_title
        try:
            os.makedirs(folder_path, exist_ok=True)
            return folder_path
        except Exception as e:
            print(f"Error creating folder: {e}")
            return None
    except Exception as e:
        print(f"Error creating folder: {e}")
        return None


def process_chapter(novel_url, chapter_number, series_folder, cookies):
    """ 处理每个章节的下载 """
    driver = create_driver(cookies)
    try:
        novel_title, content = extract_novel_content(driver, novel_url, chapter_number)
        if novel_title and content:
            save_to_word(novel_title, content, series_folder, chapter_number)
    except Exception as e:
        print(f"Error processing chapter {chapter_number}: {e}")
    finally:
        driver.quit()  # 确保每个进程关闭其浏览器实例


def download_series(series_id, status_label, progress_bar, output_label, output_path):
    cookies = [
    {"name": "first_visit_datetime_pc", "value": "2024-11-02%2016%3A36%3A23", "domain": ".pixiv.net", "path": "/"},
    {"name": "p_ab_id", "value": "7", "domain": ".pixiv.net", "path": "/"},
    {"name": "p_ab_id_2", "value": "2", "domain": ".pixiv.net", "path": "/"},
    {"name": "p_ab_d_id", "value": "1405298390", "domain": ".pixiv.net", "path": "/"},
    {"name": "yuid_b", "value": "NymXFBA", "domain": ".pixiv.net", "path": "/"},
    {"name": "PHPSESSID", "value": "84620081_AOAUJkysGvAa7wCOzTVu1jGnr5QWlyl8", "domain": ".pixiv.net", "path": "/"},
    {"name": "device_token", "value": "29ab9c0bc4d87ec35b628138a077a952", "domain": ".pixiv.net", "path": "/"},
    {"name": "privacy_policy_agreement", "value": "7", "domain": ".pixiv.net", "path": "/"},
    {"name": "c_type", "value": "25", "domain": ".pixiv.net", "path": "/"},
    {"name": "privacy_policy_notification", "value": "0", "domain": ".pixiv.net", "path": "/"},
    {"name": "a_type", "value": "0", "domain": ".pixiv.net", "path": "/"},
    {"name": "b_type", "value": "1", "domain": ".pixiv.net", "path": "/"},
    {"name": "__utmc", "value": "235335808", "domain": ".pixiv.net", "path": "/"},
    {"name": "__utmz", "value": "235335808.1731673314.3.2.utmcsr=google|utmccn=(organic)|utmcmd=organic|utmctr=(not%20provided)", "domain": ".pixiv.net", "path": "/"},
    {"name": "_gid", "value": "GA1.2.1140275552.1731673323", "domain": ".pixiv.net", "path": "/"},
    {"name": "_gcl_au", "value": "1.1.933459697.1731674003", "domain": ".pixiv.net", "path": "/"},
    {"name": "login_ever", "value": "yes", "domain": ".pixiv.net", "path": "/"},
    {"name": "FCNEC", "value": "%5B%5B%22AKsRol85ClY3AwfyMj54zcarYk00RxgcOaHHrozjnRMz5OSG9wMNsUSmXGw0NBfObg792NaeOIP13maOdW2qCHdeqzstLyzW8-3WrWNlZZ4FJKaeuBHiP7O9CGnI11fJQCny8zRdGgAf2CJbBHMFt3N_TDNPUgQ71g%3D%3D%22%5D%5D", "domain": ".pixiv.net", "path": "/"},
    {"name": "__utmv", "value": "235335808.|2=login%20ever=yes=1^3=plan=normal=1^5=gender=male=1^6=user_id=84620081=1^9=p_ab_id=7=1^10=p_ab_id_2=2=1^11=lang=zh=1", "domain": ".pixiv.net", "path": "/"},
    {"name": "__utma", "value": "235335808.1902634853.1730532986.1731682059.1731732021.6", "domain": ".pixiv.net", "path": "/"},
    {"name": "cf_clearance", "value": "bGaFX3BsTgxjfPvTG.KcsogLRRXxDmWjzhMMiNoA8po-1731732811-1.2.1.1-u___Vba3qPKSajPGYWNsiAq9uhJBEfOpS_wXr.bvHwZgYn4XcYj1NtinsFlA7weuNx0PeOqWy3cq5XbF7NRSZdSuLtqKEWa09IY7GrCkLyoj06qo6Nem3UFSEX0myY.CFzcZw8NphWe4bV90MtcoD7Ha4wZw5INzyncqyV.BZJzvAULVO9YNdkkTjAA64SqVDDjRhNJSsakRHA6Clls4zNXwsWoH1omyHXmrNC67harP8bNNUdSFWTNSpfFZi6yBrIBch1_UjJJniIpaXw5VncwBnSmPjoV_Aqd6QYSx36cFzHCZw.Gwv0FfSmA9I8yFTnv5qhM14Th28qBLq5ASjQkLcXxs7NGZ14IUWtf0iQ0GneRVKQPM4z52dwYv5Xkk", "domain": ".pixiv.net", "path": "/"},
    {"name": "_ga", "value": "GA1.2.1902634853.1730532986", "domain": ".pixiv.net", "path": "/"},
    {"name": "_ga_75BBYNYN9J", "value": "GS1.1.1731747352.9.1.1731747549.0.0.0", "domain": ".pixiv.net", "path": "/"},
    {"name": "__cf_bm", "value": "SqhvjjOCJhsBElKZXZrEcjcmU_UF.5QQMey61hTt.B0-1731762715-1.0.1.1-oAWABDh8hoL78xlGMC3AMElms8uG.x5dfm6PbwWvit1a0fNvYfJMSQ_3EiuEfMpFzNUNUxoLjz0fJxiEl0E.cwq3FH9PzJlmsyoYjWWbiwA", "domain": ".pixiv.net", "path": "/"}
    ]
    
    status_label.config(text="Downloading...")
    
    try:
        # 创建主浏览器实例
        driver = create_driver(cookies)
        series_url = f"https://www.pixiv.net/novel/series/{series_id}"
        series_folder = create_series_folder(driver, series_url, output_path)
        if not series_folder:
            status_label.config(text="创建系列文件夹失败，跳过此系列。")
            return
        
        output_label.config(text=f"Output folder: {series_folder}")
        novel_urls = extract_novel_urls(driver, series_url)
        
        total_chapters = len(novel_urls)
        progress_bar['maximum'] = total_chapters
        
        # 使用进程池并行下载章节，确保每个进程使用独立的浏览器实例
        with ProcessPoolExecutor(max_workers=5) as executor:
            futures = [
                executor.submit(process_chapter, url, i + 1, series_folder, cookies)
                for i, url in enumerate(novel_urls)
            ]
            # 等待所有任务完成
            for i, future in enumerate(futures):
                future.result()
                progress_bar['value'] = i + 1
                progress_bar.update()
        
        status_label.config(text="Completed")
    
    except Exception as e:
        status_label.config(text=f"发生错误：{e}")
    finally:
        if 'driver' in locals():
            driver.quit()


def main():
    window = tk.Tk()
    window.title("Pixiv Series Downloader")
    
    series_id_label = ttk.Label(window, text="Pixiv Series ID:")
    series_id_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")
    
    series_id_entry = ttk.Entry(window)
    series_id_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")

    output_path_label = ttk.Label(window, text="Output Path:")
    output_path_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")

    output_path_entry = ttk.Entry(window)
    output_path_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
    
    # Button to open a folder selection dialog
    def select_folder():
        folder_selected = filedialog.askdirectory()
        if folder_selected:  # If a folder is selected, update the entry
            output_path_entry.delete(0, tk.END)
            output_path_entry.insert(0, folder_selected)
    
    browse_button = ttk.Button(window, text="Browse", command=select_folder)
    browse_button.grid(row=1, column=2, padx=5, pady=5, sticky="w")
    
    status_label = ttk.Label(window, text="Ready")
    status_label.grid(row=2, column=0, columnspan=3, padx=5, pady=5)
    
    progress_bar = ttk.Progressbar(window, orient="horizontal", length=300, mode="determinate")
    progress_bar.grid(row=3, column=0, columnspan=3, padx=5, pady=5)
    
    output_label = ttk.Label(window, text="")
    output_label.grid(row=4, column=0, columnspan=3, padx=5, pady=5)
    
    def start_download():
        series_id = series_id_entry.get().strip()
        output_path = output_path_entry.get().strip()
        if not series_id:
            messagebox.showerror("Error", "Please enter a Pixiv Series ID.")
            return
        # If output_path is not provided, set a default value
        if not output_path:
            output_path = os.getcwd()  # Default to the current working directory
            messagebox.showinfo("Info", f"No output path provided. Using default path: {output_path}")
        
        download_series(series_id, status_label, progress_bar, output_label, output_path)
    
    start_button = ttk.Button(window, text="Start Download", command=start_download)
    start_button.grid(row=5, column=0, columnspan=3, padx=5, pady=10)
    
    window.mainloop()


if __name__ == "__main__":
    main()
